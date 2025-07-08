#!/usr/bin/env python3
import os
import time
import requests
import re
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from code.deepwiki2markdown import deepwiki2markdown
from code.translationmarkdown import translate_content
from code.printf import printf

def get_page_source(url):
    """
    使用 Selenium 获取指定 URL 的网页源码。

    Args:
        url (str): 要获取源码的网页 URL。

    Returns:
        str: 网页的完整 HTML 源码，如果发生错误则返回 None。
    """
    # 配置 Chrome 选项 (可选)：例如，无头模式运行
    chrome_options = Options()
    # 禁用所有日志输出
    chrome_options.add_argument('--log-level=3')  # 0=INFO, 1=WARNING, 2=ERROR, 3=FATAL
    chrome_options.add_argument('--disable-logging')  # 禁用日志
    chrome_options.add_experimental_option('excludeSwitches', ['enable-logging'])  # 排除日志开关
    # 禁用控制台输出
    chrome_options.add_argument('--disable-dev-shm-usage')
    chrome_options.add_argument('--no-sandbox')
    chrome_options.add_argument('--disable-gpu')
    chrome_options.add_argument('--headless')  # 无头模式也会减少输出
    # 指定 ChromeDriver 的路径
    # 如果您将 chromedriver 放在系统 PATH 中，则无需指定 service_executable_path
    # service = Service(executable_path='/path/to/your/chromedriver') # 替换为您的chromedriver路径
    # driver = webdriver.Chrome(service=service, options=chrome_options)
    
    # 如果 chromedriver 在 PATH 中，可以直接这样初始化
    driver = webdriver.Chrome(options=chrome_options)

    try:
        # 打开网页
        driver.get(url)

        # 等待页面加载完成 (可选，根据页面复杂度调整)
        # 您可以使用显式等待来等待某个元素出现，或者简单地等待几秒
        time.sleep(3) # 简单等待3秒，确保页面内容加载

        # 获取网页源码
        return driver.page_source

    except Exception as e:
        printf(f"An error occurred: {e}")
        return None
    finally:
        # 关闭浏览器
        if driver:
            driver.quit()

def create_table_in_doc(doc, table_lines):
    """在 Word 文档中创建表格"""
    if len(table_lines) < 2:
        return
    
    # 解析表格头
    header_line = table_lines[0]
    headers = [cell.strip() for cell in header_line.split('|') if cell.strip()]
    
    if not headers:
        return
    
    # 跳过分隔符行
    data_lines = table_lines[2:] if len(table_lines) > 2 else []
    
    # 创建表格
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # 添加表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        if i < len(hdr_cells):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    # 添加数据行
    for line in data_lines:
        cells_data = [cell.strip() for cell in line.split('|') if cell.strip()]
        if cells_data:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(cells_data):
                if i < len(row_cells):
                    row_cells[i].text = cell_data

def create_word_document(markdown_content, output_path):
    """将 Markdown 内容转换为格式化的 Word 文档"""
    doc = Document()
    
    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(12)
    
    # 解析 Markdown 内容并转换为 Word
    lines = markdown_content.split('\n')
    current_table = []
    in_code_block = False
    code_lines = []
    
    for line in lines:
        original_line = line
        line = line.strip()
        
        if not line and not in_code_block:
            continue
            
        # 处理代码块
        if line.startswith('```'):
            if in_code_block:
                # 结束代码块
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph(code_text)
                    p.style = 'No Spacing'
                    font = p.runs[0].font
                    font.name = 'Consolas'
                    font.size = Pt(10)
                code_lines = []
                in_code_block = False
            else:
                # 开始代码块
                in_code_block = True
            continue
        
        if in_code_block:
            code_lines.append(original_line)
            continue
            
        # 处理表格
        if '|' in line and line.count('|') >= 2:
            current_table.append(line)
            continue
        else:
            # 如果有积累的表格，先处理表格
            if current_table:
                create_table_in_doc(doc, current_table)
                current_table = []
        
        # 处理标题
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:], level=4)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('- '):
            # 列表项
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('* '):
            # 列表项
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('**') and line.endswith('**'):
            # 粗体段落
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
        else:
            # 普通段落
            if line:
                doc.add_paragraph(line)
    
    # 处理剩余的表格
    if current_table:
        create_table_in_doc(doc, current_table)
    
    # 保存文档
    doc.save(output_path)

def main():
    printf(">>> DeepWiki 页面 ➜  Markdown ➜  翻译 ➜  Word <<<")
    
    # 步骤 0：创建必要的目录
    directories = [
        'data/files_markdown',
        'data/files_markdown_translated',
        'data/files_word'
    ]
    for dir_path in directories:
        os.makedirs(dir_path, exist_ok=True)

    # 步骤 1：加载链接
    printf(f"\n开始提取 Excel 中的链接...")
    try:
        df = pd.read_excel('data/task.xlsx')
        col_idx = sum((ord(c.upper()) - 64) * 26 ** i for i, c in enumerate(reversed('D'))) - 1
        column_data = df.values[: , col_idx]
        url_pattern = re.compile(r'https?://[^\s]+')
        urls = {url for url in map(str, column_data) if url_pattern.match(url)}     # 这里将自动过滤重复的内容
    except Exception as e:
        printf(f"Error: {e}")
    printf(f"成功获取 {len(urls)} 个 URL！")
    
    # 步骤 2：加载网页链接并提取网页内容
    printf(f"\n开始依次提取 URL 页面内容...")
    markdown_num = 0
    for url in urls:
        try:
            printf(f"提取: {url}")
            page_source = get_page_source(url)
            # 递归处理 HTML 元素并转换为 Markdown
            markdown = deepwiki2markdown(page_source)
            
            filename = f"{url.split('/')[-1]}.md"
            md_path = os.path.join('data/files_markdown', filename)
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(markdown)
            
            printf(f"保存: {md_path}")
            markdown_num += 1
        except requests.exceptions.RequestException as e: 
            printf(f"请求错误: {e}")
        except Exception as e:
            printf(f"转换错误: {e}")
    printf(f"成功提取 {markdown_num} 个 URL 页面内容！")
    
    # 步骤 3：翻译 markdown 文件
    printf(f"\n开始翻译 Markdown 文件...")
    md_files = [f for f in os.listdir('data/files_markdown') if f.endswith('.md')]
    translated_num = 0
    for filename in enumerate(md_files):
        printf(f"翻译: {filename}")
        
        input_path = os.path.join('data/files_markdown', filename)
        output_path = os.path.join('data/files_markdown_translated', filename)
        
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 翻译内容
        # translated_content = translate_content(content)
        translated_content = 'translate_content(content)'
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(translated_content)
        
        printf(f"保存: {output_path}")
        translated_num += 1
    printf(f"成功翻译 {translated_num} 个 Markdown 文件！")

    # 步骤 4：转换为 Word 文档
    printf(f"\n开始转换为 Word 文档...")
    translated_files = [f for f in os.listdir('data/files_markdown_translated') if f.endswith('.md')]
    word_num = 0
    for filename in enumerate(translated_files):
        printf(f"转换: {filename}")
        
        input_path = os.path.join('data/files_markdown_translated', filename)
        output_path = os.path.join('data/files_word', filename.replace('.md', '.docx'))
        
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 创建 Word 文档
        try:
            create_word_document(content, output_path)
            printf(f"保存: {output_path}")
            word_num += 1
        except Exception as e:
            printf(f"转换失败: {e}")
    printf(f"成功转换 {word_num} 个 Word 文档！")

if __name__ == "__main__":
    main()
