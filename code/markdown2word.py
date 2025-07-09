import os
import re
import tempfile
import subprocess
from markdown import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from PIL import Image

def convert_markdown_to_word(input_file, output_file):
    """将Markdown文件转换为Word文档，支持Mermaid图表"""
    # 读取Markdown文件内容
    with open(input_file, 'r', encoding='utf-8') as f:
        markdown_text = f.read()
    
    # 创建临时目录存储生成的图表
    with tempfile.TemporaryDirectory() as temp_dir:
        # 处理Mermaid图表
        processed_markdown = process_mermaid_diagrams(markdown_text, temp_dir)
        
        # 将Markdown转换为HTML
        html = markdown(processed_markdown, extensions=['markdown.extensions.fenced_code'])
        
        # 创建Word文档
        doc = Document()
        
        # 解析HTML并转换为Word元素
        soup = BeautifulSoup(html, 'html.parser')
        convert_html_to_docx(soup, doc, temp_dir)
        
        # 保存Word文档
        doc.save(output_file)
        print(f"转换完成，文件已保存至: {output_file}")

def process_mermaid_diagrams(markdown_text, temp_dir):
    """处理Markdown中的Mermaid图表，生成图片并替换为图片引用"""
    # 使用正则表达式查找所有Mermaid代码块
    mermaid_pattern = re.compile(
        r'```mermaid\s*([\s\S]*?)\s*```', 
        re.MULTILINE
    )
    
    counter = 0
    processed_text = markdown_text
    
    for match in mermaid_pattern.finditer(markdown_text):
        mermaid_code = match.group(1)
        counter += 1
        
        # 生成图片文件名
        image_file = os.path.join(temp_dir, f'mermaid_{counter}.png')
        
        # 使用Mermaid CLI生成图表
        try:
            subprocess.run(
                ['mmdc', '-i', '-', '-o', image_file, '-t', 'default'],
                input=mermaid_code.encode('utf-8'),
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE
            )
            
            # 验证图片是否生成成功
            if os.path.exists(image_file):
                # 调整图片大小（如果需要）
                resize_image(image_file)
                
                # 替换Mermaid代码块为Markdown图片标记
                image_markdown = f'![Mermaid图表]({image_file})'
                processed_text = processed_text.replace(
                    match.group(0), 
                    image_markdown
                )
            else:
                print(f"警告: Mermaid图表生成失败，位置: {match.start()}")
                # 替换为错误提示
                error_markdown = f'[图表生成失败]'
                processed_text = processed_text.replace(
                    match.group(0), 
                    error_markdown
                )
        except subprocess.CalledProcessError as e:
            print(f"错误: 执行Mermaid CLI时出错: {e.stderr.decode('utf-8')}")
            # 替换为错误提示
            error_markdown = f'[图表生成失败: {e.stderr.decode("utf-8")}]'
            processed_text = processed_text.replace(
                match.group(0), 
                error_markdown
            )
        except Exception as e:
            print(f"错误: 生成Mermaid图表时出错: {e}")
            # 替换为错误提示
            error_markdown = f'[图表生成失败: {str(e)}]'
            processed_text = processed_text.replace(
                match.group(0), 
                error_markdown
            )
    
    return processed_text

def resize_image(image_path, max_width=6.0):
    """调整图片大小以适应Word文档"""
    try:
        with Image.open(image_path) as img:
            # 计算新尺寸（保持宽高比）
            width, height = img.size
            if width > 600:  # 假设Word文档宽度约为6英寸，每英寸约100像素
                ratio = max_width * 100 / width
                new_width = int(width * ratio)
                new_height = int(height * ratio)
                
                # 调整图片大小
                resized_img = img.resize((new_width, new_height), Image.LANCZOS)
                resized_img.save(image_path)
    except Exception as e:
        print(f"警告: 调整图片大小时出错: {e}")

def convert_html_to_docx(soup, doc, temp_dir):
    """将HTML内容转换为Word文档元素"""
    for element in soup.body.children if soup.body else soup.children:
        if element.name is None:
            # 处理文本节点
            if element.strip():
                doc.add_paragraph(element.strip())
            continue
            
        if element.name == 'p':
            # 处理段落
            doc.add_paragraph(element.get_text())
            
        elif element.name == 'h1':
            # 处理一级标题
            doc.add_heading(element.get_text(), level=1)
            
        elif element.name == 'h2':
            # 处理二级标题
            doc.add_heading(element.get_text(), level=2)
            
        elif element.name == 'h3':
            # 处理三级标题
            doc.add_heading(element.get_text(), level=3)
            
        elif element.name == 'h4':
            # 处理四级标题
            doc.add_heading(element.get_text(), level=4)
            
        elif element.name == 'ul':
            # 处理无序列表
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Bullet')
                
        elif element.name == 'ol':
            # 处理有序列表
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Number')
                
        elif element.name == 'img':
            # 处理图片
            src = element.get('src')
            if src and os.path.exists(src) and os.path.basename(src).startswith('mermaid_'):
                # 添加Mermaid生成的图片
                doc.add_picture(src, width=Inches(6))
            else:
                # 添加普通图片（如果有）
                doc.add_paragraph(f"[图片: {src}]")
                
        elif element.name == 'pre':
            # 处理代码块
            if element.code:
                code_text = element.code.get_text()
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(code_text)
                run.font.name = 'Consolas'  # 设置代码字体
        # 可以根据需要添加更多元素类型的处理

# if __name__ == "__main__":
#     import argparse
    
#     parser = argparse.ArgumentParser(description='将Markdown文件转换为Word文档，支持Mermaid图表')
#     parser.add_argument('input', help='输入的Markdown文件路径')
#     parser.add_argument('-o', '--output', help='输出的Word文件路径', default=None)
    
#     args = parser.parse_args()
    
#     # 确定输出文件路径
#     if args.output:
#         output_file = args.output
#     else:
#         base_name = os.path.splitext(args.input)[0]
#         output_file = f"{base_name}.docx"
    
#     # 执行转换
#     convert_markdown_to_word(args.input, output_file)    